import openai
import os
from pydub import AudioSegment
import re
from io import BytesIO
from docx import Document

# Function to split audio into chunks
def format_audio(audio_file_path, chunk_duration_ms=360000):  # 6 minutes chunks
    audio = AudioSegment.from_file(audio_file_path)
    duration_ms = len(audio)
    chunks = []

    for start_time in range(0, duration_ms, chunk_duration_ms):
        end_time = min(start_time + chunk_duration_ms, duration_ms)
        chunk = audio[start_time:end_time]
        chunks.append(chunk)

    return chunks

# Function for formatting transcript
def format_transcription(transcript):
    prompt = (
        "以下の文字起こしテキストを読みやすく整形してください。内容はそのままに、段落分けや句読点の追加などを行ってください。\n\n"
        f"{transcript}"
    )

    response = openai.chat.completions.create(
        messages=[
            {"role": "system", "content": "あなたは優秀な日本語のエディターです。"},
            {"role": "user", "content": prompt}
        ],
        model="gpt-4o-mini",
        max_tokens=16384,
        temperature=0
    )

    return response.choices[0].message.content

# Function to transcribe the audio chunks
def transcribe_audio(chunks):
    client = openai.OpenAI()
    full_transcript = ""

    for i, chunk in enumerate(chunks):
        # Save chunk to a temporary file
        chunk_path = f"temp_chunk_{i}.mp3"
        chunk.export(chunk_path, format="mp3")

        # Transcribe audio chunk
        with open(chunk_path, 'rb') as audio_file:
            transcript = client.audio.transcriptions.create(
                file=audio_file,
                model="whisper-1",
                language="ja",
                response_format="text",
            )

        formatted_transcript = format_transcription(transcript)  # Format transcript
        full_transcript += formatted_transcript + "\n" # Append chunks
        os.remove(chunk_path) # Remove the temporary chunk file

    return full_transcript

# Function to save transcription as a .docx file
def save_transcription_to_docx(transcription_text):
    doc = Document()
    doc.add_heading('音声文字起こし結果', 0)

    # Add each paragraph from the transcription
    paragraphs = transcription_text.split("\n")
    for para in paragraphs:
        doc.add_paragraph(para)

    # Save the document in memory
    doc_io = BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io

def create_outline(transcript):
    """Creates an outline from a transcript."""
    prompt = (
        "All operations and responses must be conducted in Japanese . Only provide the requested output, any conversation between the LLM and the user is unnecessary.\n"
        "提供されたトランスクリプトをもとに、レポートの詳細なアウトライン（見出しのサブセクションまで）を作成してください。この際、トランスクリプトに含まれない内容は入れないでください。\n"
        "レポートの構成にイントロとまとめの章は不要であり、本論のみを考えてください。最大4つの章までに収め、各章「第１章」のようにタイトルをつけてください。\n\n"
        f"トランスクリプト:\n{transcript}"
    )
    response = openai.chat.completions.create(
        messages=[
            {"role": "system", "content": "あなたはプロのライターです。"},
            {"role": "user", "content": prompt}
        ],
        model="gpt-4o-mini",
        max_tokens=16384,
        temperature=0
    )
    return response.choices[0].message.content

def write_chapters(transcript, outline, num_chapters):
    """Writes chapters based on the transcript and outline."""
    written_chapters = []
    for i in range(num_chapters):
        n = i + 1
        part_name = f"第{n}章"
        prompt = (
            "All operations and responses must be conducted in Japanese. Only provide the requested output, any conversation between the gpt and the user is unnecessary. \n"
            f"提供したトランスクリプトの内容をもとに、アウトラインの{part_name}を詳細に執筆してください。トランスクリプトの文量に合わせて執筆する分量を調整してください。また、トランスクリプトに含まれない内容は書かないでください。\n"
            f"アウトライン:\n{outline}\n\n"
            f"トランスクリプト:\n{transcript}\n\n" #要修正：執筆完了した部分を削除しトークン数を減らしたい
        )
        response = openai.chat.completions.create(
            messages=[
                {"role": "system", "content": "あなたはプロのライターです。"},
                {"role": "user", "content": prompt}
            ],
            model="gpt-4o-mini",
            max_tokens=16384,
            temperature=0
        )
        chapter_content = response.choices[0].message.content
        written_chapters.append(chapter_content)
    return written_chapters

def save_report_to_docx(title, outline, written_chapters):
    """Saves the final report as a .docx file."""
    doc = Document()
    doc.add_heading(f"{title}の概要", 0)
    doc.add_heading("アウトライン", level=1)
    doc.add_paragraph(outline)

    for i, chapter in enumerate(written_chapters):
        n = i + 1
        part_name = f"第{n}章"
        doc.add_heading(part_name, level=1)
        doc.add_paragraph(chapter)
    
    doc_io = BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io
